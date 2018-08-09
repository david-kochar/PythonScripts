# -*- coding: utf-8 -*-
"""
Created on Tue Aug  7 20:11:45 2018

@author: dkochar
"""
# note that the "docx" library installation was previously performed using 
# "conda install -c conda-forge python-docx" in Anaconda CMD. You will need to
# download the docx library zip file from https://pypi.org/project/docx/#files

#import libraries
import pandas as pnds
import io
import requests
import docx

# create a dataframe from a URL-based csv. We will use the ToothGrowth dataset
csvurl = "https://vincentarelbundock.github.io/Rdatasets/csv/datasets/ToothGrowth.csv"
req    = requests.get(csvurl).content
df     = pnds.DataFrame(pnds.read_csv(io.StringIO(req.decode('utf-8'))))

# create Microsoft Word Document
from docx import Document

document = Document()
document.save('embedtabletest.docx')

# open the previosuly saved document
doc = docx.Document('./embedtabletest.docx')

# create a table in the Word Document with the same dimensions as the dataframe
t = doc.add_table(df.shape[0]+1, df.shape[1])

for j in range(df.shape[-1]):
    t.cell(0,j).text = df.columns[j]

#format table style to be a grid    
t.style = 'TableGrid'

# populate the table with the dataframe
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
        t.cell(i+1,j).text = str(df.values[i,j])

#get table dimensions to use for a subsequent step
row = t.rows[0]
column = t.columns[0]
trows = len(column.cells) * 3 #multiplier is a work-around for xml parsing limitation
tcols = len(row.cells)

# Highlight all cells limegreen (RGB 32CD32) if cell contains text "0.5" 
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

for i in range(tcols):
    for j in range(trows):
       if t.cell(i,j).text == '0.5':
           t.cell(i,j)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="32CD32"/>'.format(nsdecls('w'))))            
    
# save the modified document
doc.save('./embedtabletest.docx')
